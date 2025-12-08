import io
from docx import Document
import pdfplumber

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    print(text_parts)
    return "\n".join(text_parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts all available text from a DOCX file, including paragraphs and table cells.
    
    Args:
        file_bytes (bytes): The raw binary content of a DOCX file.

    Returns:
        str: All extracted text, separated by newlines.
    """
    document = Document(io.BytesIO(file_bytes))
    text_parts = []

    # Extract all paragraph text
    for para in document.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    print(text_parts)
    return "\n".join(text_parts)


def extract_text(file_bytes: bytes, name: str) -> str:
    ext = name.lower().split('.')[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension")
