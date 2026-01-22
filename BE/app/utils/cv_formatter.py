"""CV Formatter - Generate editable and professional CV documents (DOCX and PDF)."""
from typing import Dict, List, Optional, Any
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class CVFormatter:
    """Format CV data into professional DOCX documents."""
    
    # Color scheme
    PRIMARY_COLOR = RGBColor(31, 78, 121)      # Dark blue
    ACCENT_COLOR = RGBColor(192, 0, 0)         # Red
    TEXT_COLOR = RGBColor(64, 64, 64)          # Dark gray
    LIGHT_TEXT = RGBColor(128, 128, 128)       # Light gray
    
    # Font settings
    NAME_FONT_SIZE = 18
    SECTION_FONT_SIZE = 12
    BODY_FONT_SIZE = 11
    SMALL_FONT_SIZE = 10
    
    def __init__(self):
        """Initialize CV formatter."""
        self.doc = None
    
    def create_cv(self, cv_data: Dict[str, Any]) -> Document:
        """Create a professional DOCX CV document.
        
        Args:
            cv_data: Dictionary containing CV information with structure:
                {
                    'name': str,
                    'email': str,
                    'phone': str (optional),
                    'location': str (optional),
                    'summary': str (optional),
                    'experience': [
                        {'company': str, 'role': str, 'dates': str, 'description': str},
                        ...
                    ],
                    'education': [
                        {'degree': str, 'institution': str, 'year': str},
                        ...
                    ],
                    'skills': [str, ...],
                    'certifications': [str, ...] (optional),
                    'projects': [
                        {'project': str, 'description': str},
                        ...
                    ] (optional),
                }
        
        Returns:
            Document: The created Word document
        """
        self.doc = Document()
        self._set_document_margins()
        
        # Add header
        self._add_header(cv_data)
        
        # Add summary if present
        if cv_data.get('summary'):
            self._add_summary(cv_data['summary'])
        
        # Add sections
        if cv_data.get('experience'):
            self._add_experience_section(cv_data['experience'])
        
        if cv_data.get('education'):
            self._add_education_section(cv_data['education'])
        
        if cv_data.get('skills'):
            self._add_skills_section(cv_data['skills'])
        
        if cv_data.get('certifications'):
            self._add_certifications_section(cv_data['certifications'])
        
        if cv_data.get('projects'):
            self._add_projects_section(cv_data['projects'])
        
        return self.doc
    
    def _set_document_margins(self, top=1, bottom=1, left=0.75, right=0.75):
        """Set document margins in inches."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
    
    def _add_header(self, cv_data: Dict[str, str]):
        """Add CV header with name and contact information."""
        # Name
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(cv_data.get('name', 'Your Name'))
        name_run.font.size = Pt(self.NAME_FONT_SIZE)
        name_run.font.bold = True
        name_run.font.color.rgb = self.PRIMARY_COLOR
        
        # Contact information
        contact_info = []
        if cv_data.get('email'):
            contact_info.append(cv_data['email'])
        if cv_data.get('phone'):
            contact_info.append(cv_data['phone'])
        if cv_data.get('location'):
            contact_info.append(cv_data['location'])
        
        if contact_info:
            contact_para = self.doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(self.SMALL_FONT_SIZE)
            contact_run.font.color.rgb = self.LIGHT_TEXT
        
        # Add horizontal line
        self._add_horizontal_line()
        
        # Add spacing
        self.doc.add_paragraph()
    
    def _add_summary(self, summary: str):
        """Add professional summary section."""
        self._add_section_title('PROFESSIONAL SUMMARY')
        summary_para = self.doc.add_paragraph(summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._format_body_text(summary_para)
        self.doc.add_paragraph()
    
    def _add_experience_section(self, experience: List[Dict[str, str]]):
        """Add work experience section."""
        self._add_section_title('PROFESSIONAL EXPERIENCE')
        
        for i, job in enumerate(experience):
            # Job title and company
            job_para = self.doc.add_paragraph()
            
            # Company and role
            company_run = job_para.add_run(f"{job.get('company', 'Company')}")
            company_run.font.bold = True
            company_run.font.size = Pt(self.BODY_FONT_SIZE)
            company_run.font.color.rgb = self.TEXT_COLOR
            
            # Role in separate line if available
            if job.get('role'):
                job_para.add_run('\n')
                role_run = job_para.add_run(f"{job.get('role', '')}")
                role_run.font.italic = True
                role_run.font.size = Pt(self.BODY_FONT_SIZE)
                role_run.font.color.rgb = self.TEXT_COLOR
            
            # Dates
            if job.get('dates'):
                dates_para = self.doc.add_paragraph(job.get('dates', ''))
                dates_para.paragraph_format.left_indent = Inches(0.25)
                self._format_subtext(dates_para)
            
            # Description
            if job.get('description'):
                desc_para = self.doc.add_paragraph(job.get('description', ''), style='List Bullet')
                desc_para.paragraph_format.left_indent = Inches(0.5)
                self._format_body_text(desc_para)
            
            # Spacing between entries
            if i < len(experience) - 1:
                self.doc.add_paragraph()
        
        self.doc.add_paragraph()
    
    def _add_education_section(self, education: List[Dict[str, str]]):
        """Add education section."""
        self._add_section_title('EDUCATION')
        
        for i, edu in enumerate(education):
            # Degree and institution
            edu_para = self.doc.add_paragraph()
            
            degree_run = edu_para.add_run(f"{edu.get('degree', 'Degree')}")
            degree_run.font.bold = True
            degree_run.font.size = Pt(self.BODY_FONT_SIZE)
            degree_run.font.color.rgb = self.TEXT_COLOR
            
            if edu.get('institution'):
                edu_para.add_run('\n')
                inst_run = edu_para.add_run(f"{edu.get('institution', '')}")
                inst_run.font.size = Pt(self.BODY_FONT_SIZE)
                inst_run.font.color.rgb = self.TEXT_COLOR
            
            # Year
            if edu.get('year'):
                year_para = self.doc.add_paragraph(edu.get('year', ''))
                year_para.paragraph_format.left_indent = Inches(0.25)
                self._format_subtext(year_para)
            
            # Spacing between entries
            if i < len(education) - 1:
                self.doc.add_paragraph()
        
        self.doc.add_paragraph()
    
    def _add_skills_section(self, skills: List[str]):
        """Add skills section."""
        self._add_section_title('TECHNICAL SKILLS')
        
        # Group skills for better presentation
        skills_str = ', '.join(skills)
        skills_para = self.doc.add_paragraph(skills_str)
        skills_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._format_body_text(skills_para)
        
        self.doc.add_paragraph()
    
    def _add_certifications_section(self, certifications: List[str]):
        """Add certifications section."""
        self._add_section_title('CERTIFICATIONS & LICENSES')
        
        for cert in certifications:
            cert_para = self.doc.add_paragraph(cert, style='List Bullet')
            self._format_body_text(cert_para)
        
        self.doc.add_paragraph()
    
    def _add_projects_section(self, projects: List[Dict[str, str]]):
        """Add projects section."""
        self._add_section_title('PROJECTS')
        
        for i, project in enumerate(projects):
            # Project name
            proj_para = self.doc.add_paragraph()
            proj_run = proj_para.add_run(f"{project.get('project', 'Project')}")
            proj_run.font.bold = True
            proj_run.font.size = Pt(self.BODY_FONT_SIZE)
            proj_run.font.color.rgb = self.TEXT_COLOR
            
            # Description
            if project.get('description'):
                desc_para = self.doc.add_paragraph(project.get('description', ''), style='List Bullet')
                desc_para.paragraph_format.left_indent = Inches(0.5)
                self._format_body_text(desc_para)
            
            # Spacing between entries
            if i < len(projects) - 1:
                self.doc.add_paragraph()
        
        self.doc.add_paragraph()
    
    def _add_section_title(self, title: str):
        """Add a section title with formatting."""
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(self.SECTION_FONT_SIZE)
        title_run.font.bold = True
        title_run.font.color.rgb = self.PRIMARY_COLOR
        
        # Add underline
        pPr = title_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # Border size
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self._rgb_to_hex(self.PRIMARY_COLOR))
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        self.doc.add_paragraph()
    
    def _add_horizontal_line(self):
        """Add a horizontal line separator."""
        para = self.doc.add_paragraph()
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self._rgb_to_hex(self.PRIMARY_COLOR))
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _format_body_text(self, paragraph):
        """Apply formatting to body text."""
        for run in paragraph.runs:
            run.font.size = Pt(self.BODY_FONT_SIZE)
            run.font.color.rgb = self.TEXT_COLOR
    
    def _format_subtext(self, paragraph):
        """Apply formatting to subtext (dates, etc)."""
        for run in paragraph.runs:
            run.font.size = Pt(self.SMALL_FONT_SIZE)
            run.font.italic = True
            run.font.color.rgb = self.LIGHT_TEXT
    
    @staticmethod
    def _rgb_to_hex(rgb: RGBColor) -> str:
        """Convert RGBColor to hex string."""
        return f'{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
    
    def save_to_file(self, filepath: str) -> str:
        """Save document to file.
        
        Args:
            filepath: Path where to save the DOCX file
            
        Returns:
            str: The filepath where document was saved
        """
        if self.doc:
            self.doc.save(filepath)
            return filepath
        raise ValueError("No document to save. Create a CV first.")
    
    def save_to_bytes(self) -> BytesIO:
        """Save document to BytesIO object.
        
        Returns:
            BytesIO: Document as bytes
        """
        if self.doc:
            byte_stream = BytesIO()
            self.doc.save(byte_stream)
            byte_stream.seek(0)
            return byte_stream
        raise ValueError("No document to save. Create a CV first.")


def format_cv_from_parsed_sections(cv_text: str, cv_data: Optional[Dict[str, Any]] = None) -> Document:
    """Format a CV from parsed sections or raw data.
    
    Args:
        cv_text: Raw CV text (for fallback)
        cv_data: Structured CV data dictionary
        
    Returns:
        Document: Formatted Word document
    """
    formatter = CVFormatter()
    
    # Use provided data or create from text
    if cv_data:
        return formatter.create_cv(cv_data)
    else:
        # Fallback: create simple formatted CV from text
        fallback_data = {
            'name': 'Formatted CV',
            'email': 'contact@example.com',
            'summary': cv_text[:500],  # Use first 500 chars as summary
        }
        return formatter.create_cv(fallback_data)


def format_cv_professional(parsed_sections: Dict, original_cv_text: str) -> Document:
    """Format CV with professional styling from parsed sections.
    
    Args:
        parsed_sections: Dictionary of parsed CV sections
        original_cv_text: Original CV text for reference
        
    Returns:
        Document: Professionally formatted Word document
    """
    formatter = CVFormatter()
    
    # Build structured data from parsed sections
    cv_data = {
        'name': parsed_sections.get('contact', {}).get('title', 'Your Name'),
        'email': parsed_sections.get('contact', {}).get('items', [{}])[0].get('label', ''),
        'experience': [
            item for item in parsed_sections.get('experience', {}).get('items', [])
        ] if parsed_sections.get('experience') else [],
        'education': [
            item for item in parsed_sections.get('education', {}).get('items', [])
        ] if parsed_sections.get('education') else [],
        'skills': [
            item.get('skill', item.get('label', '')) 
            for item in parsed_sections.get('skills', {}).get('items', [])
        ] if parsed_sections.get('skills') else [],
        'certifications': [
            item.get('label', '') 
            for item in parsed_sections.get('certifications', {}).get('items', [])
        ] if parsed_sections.get('certifications') else [],
        'projects': [
            item for item in parsed_sections.get('projects', {}).get('items', [])
        ] if parsed_sections.get('projects') else [],
    }
    
    return formatter.create_cv(cv_data)
